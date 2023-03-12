from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import statistics


class WordHandler:
    def __init__(self, filePath, games):
        self.filePath = filePath
        self.games = games
        self.medianOfAllGames = 0
        self.SDofAllGames = 0
        self.medianOfWhiteGames = 0
        self.SDofWhiteGames = 0
        self.medianOfBlackGames = 0
        self.SDofBlackGames = 0

    def getGames(self):
        return self.games

    def isWhiteGame(self, game):
        return game.getWhitePlayer().getName().startswith("Stockfish")

    def isBlackGame(self, game):
        return game.getBlackPlayer().getName().startswith("Stockfish")

    def isStockfishWin(self, game):
        return (self.isWhiteGame(game) and game.getResult() == "1-0") or (self.isBlackGame(game) and game.getResult() == "0-1")

    def isStockfishLoss(self, game):
        return (self.isWhiteGame(game) and game.getResult() == "0-1") or (self.isBlackGame(game) and game.getResult() == "1-0")

    def filterOpeningsPlayedNTimes(self, n=1):
        if n < 1:
            raise ValueError("n must be a positive integer.")
        openings = dict()
        for game in self.getGames():
            if game.getOpening() not in openings.keys():
                openings[game.getOpening()] = [0, 0, 0]
            result = game.getResult()
            if result == "1-0":
                openings[game.getOpening()][0] += 1
            elif result == "1/2-1/2":
                openings[game.getOpening()][1] += 1
            elif result == "0-1":
                openings[game.getOpening()][2] += 1

        # Checking that there exists an opening with n amount of games
        if not bool(openings):
            raise ValueError("Choose a lower amount of games.")

        popularOpenings = dict()
        for key, value in openings.items():
            if sum(value) >= n:
                popularOpenings[key] = value

        if not bool(popularOpenings):
            raise ValueError("No openings played {n} times. Choose a lower amount of games.")

        return popularOpenings

    def getGamesWithOpening(self, games, opening):
        relevantGames = []
        for game in games:
            if game.getOpening() == opening:
                relevantGames.append(game)
        if relevantGames == []:
            raise ValueError("No games with that opening")
        return relevantGames
    
    def createDocument(self, listOfPNGs, openings=[], minAmountOfGames=1):
        document = Document()
        document.add_heading('Chess game', 0)
        self.addResultTables(document, self.getGames())
        self.addMedianSDTables(document)
        document.add_picture('Datafiles\proportionResultsPlot.png')
        document.add_picture('Datafiles\proportionStockfishPlot.png')
        if openings:
            for opening in openings:
                relevantGames = self.getGamesWithOpening(self.getGames(), opening)
                self.addResultTables(document, relevantGames, opening)
        if minAmountOfGames:
            self.addOpeningsWithNGamesTable(document, minAmountOfGames)
        for png in listOfPNGs:
            document.add_picture(png, width = Inches(6), height = Inches(6))
        document.save(self.filePath)
   
    
    # Creates a table with the amount of games won, drawn and lost by Stockfish
    def addResultTables(self, document, games, opening=0):
        whiteWon = whiteDrawn = whiteLost = blackWon = blackDrawn = blackLost = 0
        for game in games:
            if game.getWhitePlayer().getName().startswith("Stockfish"):
                if game.getResult() == "1-0":
                    whiteWon += 1
                elif game.getResult() == "0-1":
                    whiteLost += 1
                elif game.getResult() == "1/2-1/2":
                    whiteDrawn += 1
            elif game.getBlackPlayer().getName().startswith("Stockfish"):
                if game.getResult() == "0-1":
                    blackWon += 1
                elif game.getResult() == "1-0":
                    blackLost += 1
                elif game.getResult() == "1/2-1/2":
                    blackDrawn += 1

        gamesWon = whiteWon + blackWon
        gamesDrawn = whiteDrawn + blackDrawn
        gamesLost = whiteLost + blackLost

        allGamesRecords = (
            ("Won", "Stockfish all games", [gamesWon, whiteWon, blackWon]),
            ("Drawn", "Stockfish white games", [
             gamesDrawn, whiteDrawn, blackDrawn]),
            ("Lost", "Stockfish black games", [
             gamesLost, whiteLost, blackLost])
        )
        if opening:
            document.add_heading(f"Games played with the {opening} opening", 2)

        for i in range(len(allGamesRecords)):
            document.add_heading(allGamesRecords[i][1], 3)
            table = document.add_table(rows=1, cols=2)
            table.style = "Medium Shading 1"
            header = table.rows[0].cells
            header[0].text = "Result"
            header[1].text = "Amount"
            for type, _, value in allGamesRecords:
                row_cells = table.add_row().cells
                row_cells[0].text = type
                row_cells[1].text = str(value[i])

    # Creates a table with statistics of each opening played at least n times
    def addOpeningsWithNGamesTable(self, document, amountOfGames):
        document.add_heading(
            f"All openings with at least {amountOfGames} games played.")
        openingsTable = document.add_table(rows=1, cols=4)
        openingsTable.style = "Medium Shading 1"
        openingsHeader = openingsTable.rows[0].cells
        openingsHeader[0].text = "Opening"
        openingsHeader[1].text = "White won"
        openingsHeader[2].text = "Draw"
        openingsHeader[3].text = "Black won"
        openings = self.filterOpeningsPlayedNTimes(amountOfGames)
        for key, value in openings.items():
            row = openingsTable.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value[0])
            row[2].text = str(value[1])
            row[3].text = str(value[2])

    # Creates a table with the median and standard deviation of the length of games
    def addMedianSDTables(self, document):
        document.add_heading(
            "Median and standard deviation for length of games", 3)
        medianSDTable = document.add_table(rows=1, cols=2)
        medianSDTable.style = "Medium Shading 1"
        tableHeader = medianSDTable.rows[0].cells
        tableHeader[0].text = "Type"
        tableHeader[1].text = "Value"

        records = (
            ("Median all games", self.medianOfAllGames),
            ("SD all games", self.SDofAllGames),
            ("Median white games", self.medianOfWhiteGames),
            ("SD white games", self.SDofWhiteGames),
            ("Median black games", self.medianOfBlackGames),
            ("SD black games", self.SDofBlackGames)
        )

        for type, value in records:
            row_cells = medianSDTable.add_row().cells
            row_cells[0].text = type
            row_cells[1].text = str(round(value, 2))

    def createResultGraph(self):
        longestGameCount = int(self.getGames()[0].getPlyCount())
        longestWhiteGameCount = longestBlackGameCount = 0
        movesAllGames = []
        movesWhiteGames = []
        movesBlackGames = []
        gamesStockfishWhite = []
        gamesStockfishBlack = []

        for game in self.getGames():
            # Find stockfish white and black games
            if self.isWhiteGame(game):
                gamesStockfishWhite.append(game)
                movesWhiteGames.append(int(game.getPlyCount()))
                if int(game.getPlyCount()) > longestWhiteGameCount:
                    longestWhiteGameCount = int(game.getPlyCount())
            elif self.isBlackGame(game):
                gamesStockfishBlack.append(game)
                movesBlackGames.append(int(game.getPlyCount()))
                if int(game.getPlyCount()) > longestBlackGameCount:
                    longestBlackGameCount = int(game.getPlyCount())

            # Find longest game
            movesAllGames.append(int(game.getPlyCount()))
            if int(game.getPlyCount()) > longestGameCount:
                longestGameCount = int(game.getPlyCount())

        nrOfGames = len(self.getGames())
        nrOfWhiteGames = len(gamesStockfishWhite)
        nrOfBlackGames = len(gamesStockfishBlack)

        # Find median and SD of all games, stockfish games with white and stockfish games with black
        self.medianOfAllGames = statistics.median(movesAllGames)
        self.SDofAllGames = statistics.pstdev(movesAllGames)
        if movesWhiteGames:
            self.medianOfWhiteGames = statistics.median(movesWhiteGames)
            self.SDofWhiteGames = statistics.pstdev(movesWhiteGames)
        if movesBlackGames:
            self.medianOfBlackGames = statistics.median(movesBlackGames)
            self.SDofBlackGames = statistics.pstdev(movesBlackGames)

        moveList = list(range(1, int(longestGameCount) + 2))
        whiteMoveList = list(range(1, int(longestWhiteGameCount) + 2))
        blackMoveList = list(range(1, int(longestBlackGameCount) + 2))
        ongoingGamesOnMove = [0] * (int(longestGameCount) + 1)
        ongoingWhiteGamesOnMove = [0] * (int(longestWhiteGameCount) + 1)
        ongoingBlackGamesOnMove = [0] * (int(longestBlackGameCount) + 1)
        for game in self.getGames():
            gameLength = int(game.getPlyCount())
            if self.isWhiteGame(game):
                for i in range(gameLength):
                    ongoingWhiteGamesOnMove[i] += 1
            elif self.isBlackGame(game):
                for i in range(gameLength):
                    ongoingBlackGamesOnMove[i] += 1
            for i in range(gameLength):
                ongoingGamesOnMove[i] += 1

        proportionOfGames = [] 
        proportionOfWhiteGames = [] 
        proportionOfBlackGames = []
        for i in range(len(ongoingGamesOnMove)):
            proportionOfGames.append(
                round(ongoingGamesOnMove[i]*100/nrOfGames, 2))

        for i in range(len(ongoingWhiteGamesOnMove)):
            proportionOfWhiteGames.append(
                round(ongoingWhiteGamesOnMove[i]*100/nrOfWhiteGames, 2))

        for i in range(len(ongoingBlackGamesOnMove)):
            proportionOfBlackGames.append(
                round(ongoingBlackGamesOnMove[i]*100/nrOfBlackGames, 2))

        plt.figure(1)
        plt.plot(moveList, proportionOfGames, 'b', label="All games")
        plt.plot(whiteMoveList, proportionOfWhiteGames,
                 'r', label="White games")
        plt.plot(blackMoveList, proportionOfBlackGames,
                 'g', label="Black games")
        plt.xlim(0, longestGameCount+2)
        plt.ylim(0, 102)
        plt.xlabel("Number of moves")
        plt.ylabel("Percentage of games ongoing")
        plt.suptitle("Proportion of games still ongoing after n moves")
        plt.legend(loc="upper right")
        plt.savefig("Datafiles\proportionResultsPlot.png")

    # Create a graph of the results of the games played by Stockfish
    def createSFResultGraph(self):
        longestGameCount = int(self.getGames()[0].getPlyCount())
        wonGames = []
        lostGames = []
        longestWonGameCount = longestLostGameCount = 0

        for game in self.getGames():
            if self.isStockfishWin(game):
                wonGames.append(game)
                if int(game.getPlyCount()) > longestWonGameCount:
                    longestWonGameCount = int(game.getPlyCount())
            elif self.isStockfishLoss(game):
                lostGames.append(game)
                if int(game.getPlyCount()) > longestLostGameCount:
                    longestLostGameCount = int(game.getPlyCount())

        longestGameCount = max(longestWonGameCount, longestLostGameCount)
        nrOfWonGames = len(wonGames)
        nrOfLostGames = len(lostGames)

        wonMoveList = list(range(1, int(longestWonGameCount) + 2))
        lostMoveList = list(range(1, int(longestLostGameCount) + 2))

        ongoingWonGamesOnMove = [0] * (int(longestWonGameCount) + 1)
        ongoingLostGamesOnMove = [0] * (int(longestLostGameCount) + 1)

        for game in self.getGames():
            gameLength = int(game.getPlyCount())
            if self.isStockfishWin(game):
                for i in range(gameLength):
                    ongoingWonGamesOnMove[i] += 1
            elif self.isStockfishLoss(game):
                for i in range(gameLength):
                    ongoingLostGamesOnMove[i] += 1

        proportionOfWonGames = []
        proportionOfLostGames = []

        for i in range(len(ongoingWonGamesOnMove)):
            proportionOfWonGames.append(
                round(ongoingWonGamesOnMove[i]*100/nrOfWonGames))

        for i in range(len(ongoingLostGamesOnMove)):
            proportionOfLostGames.append(
                round(ongoingLostGamesOnMove[i]*100/nrOfLostGames))

        plt.figure(2)
        plt.plot(wonMoveList, proportionOfWonGames, 'g', label="Games won")
        plt.plot(lostMoveList, proportionOfLostGames, 'r', label="Games lost")
        plt.xlim(0, longestGameCount+2)
        plt.ylim(0, 102)
        plt.xlabel("Number of moves")
        plt.ylabel("Percentage of games ongoing")
        plt.suptitle(
            "Proportion of Stockfish games still ongoing after n moves")
        plt.legend(loc="upper right")
        plt.savefig("Datafiles\proportionStockfishPlot.png")

    def createGraphs(self):
        self.createResultGraph()
        self.createSFResultGraph()